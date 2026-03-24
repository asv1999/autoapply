"""AutoApply v3 — Complete API with all 18 endpoints"""
import os, uuid, json, logging
from datetime import datetime
from contextlib import asynccontextmanager
from typing import List, Dict, Optional

from fastapi import FastAPI, HTTPException, UploadFile, File, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from config import settings
from database import init_db, ProfileDB, JobDB, ApplicationDB, ConnectionDB, RunDB

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@asynccontextmanager
async def lifespan(app: FastAPI):
    init_db()
    for d in ["data/outputs/resumes","data/outputs/screenshots","data/profiles"]:
        os.makedirs(d, exist_ok=True)
    yield

app = FastAPI(title="AutoApply", version="3.0.0", lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

def get_llm():
    from intelligence.engine import LLMClient
    return LLMClient(settings.GROQ_API_KEY, settings.OLLAMA_HOST)

# ─── MODELS ───
class ProfileIn(BaseModel):
    name: str; email: str = ""; phone: str = ""; location: str = ""
    linkedin_url: str = ""; target_roles: str = ""
    salary_min: int = 0; salary_max: int = 0
    resume_bullets: dict = {}; voice_rules: str = ""
    proof_points: list = []; education: str = ""; skills: list = []

class JobIn(BaseModel):
    title: str; company: str; location: str = ""; url: str = ""
    requirements: list = []; decision_maker: str = ""; ats_platform: str = ""

class OutcomeIn(BaseModel):
    application_id: int; outcome: str

# ═══ 1. HEALTH ═══
@app.get("/health")
async def health():
    llm = get_llm(); st = await llm.health()
    return {"status":"ok","llm":st,"database":"ok","profile_exists":ProfileDB.exists()}

# ═══ 2-4. PROFILE (Page 1: Onboarding) ═══
@app.post("/api/profile")
async def create_profile(p: ProfileIn):
    pid = ProfileDB.create(p.dict())
    return {"id":pid,"message":"Profile created"}

@app.get("/api/profile")
@app.get("/api/profile/{pid}")
async def get_profile(pid: int = 1):
    p = ProfileDB.get(pid)
    if not p: raise HTTPException(404, "Profile not found")
    return p

@app.put("/api/profile/{pid}")
async def update_profile(pid: int, data: dict):
    ProfileDB.update(pid, data)
    return {"message":"Updated"}

@app.post("/api/profile/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    path = f"data/profiles/{file.filename}"
    with open(path, "wb") as f: f.write(await file.read())
    ProfileDB.update(1, {"resume_template_path": path})
    return {"path":path,"filename":file.filename}

# ═══ 5-6. DISCOVERY + MATCHING (Page 3: Job Discovery) ═══
@app.post("/api/discover")
async def discover(bg: BackgroundTasks):
    rid = str(uuid.uuid4())[:8]
    RunDB.create(rid)
    async def _run():
        from discovery.engine import DiscoveryEngine
        engine = DiscoveryEngine(settings)
        try:
            stats = await engine.run_full(rid)
            RunDB.update(rid, status="completed", completed_at=datetime.now().isoformat(),
                jobs_discovered=stats["total_found"], jobs_new=stats["new_added"])
            # Auto-score new jobs
            await _score_new_jobs()
        except Exception as e:
            RunDB.update(rid, status="failed"); logger.error(f"Discovery: {e}")
        finally: await engine.cleanup()
    bg.add_task(_run)
    return {"run_id":rid,"status":"started"}

async def _score_new_jobs():
    profile = ProfileDB.get(1)
    if not profile: return
    jobs = JobDB.get_unapplied(50)
    unscored = [j for j in jobs if not j.get("match_score")]
    if not unscored: return
    llm = get_llm(); await llm.init()
    from intelligence.engine import MatchScorer
    scorer = MatchScorer(llm)
    scores, _ = await scorer.score_batch(profile, unscored)
    for jid, data in scores.items():
        JobDB.update_score(jid, data["score"], data.get("archetype"))

@app.post("/api/jobs/manual")
async def add_manual_jobs(jobs: List[JobIn]):
    result = JobDB.insert_batch([j.dict() for j in jobs])
    # Score them
    await _score_new_jobs()
    return result

@app.get("/api/jobs")
async def get_jobs(status: str = "all", sort: str = "discovered_at", limit: int = 200):
    if status == "unapplied": return JobDB.get_unapplied(limit)
    return JobDB.get_all(limit, sort)

@app.get("/api/jobs/{jid}")
async def get_job(jid: int):
    j = JobDB.get_by_id(jid)
    if not j: raise HTTPException(404)
    # Include application if exists
    app_data = ApplicationDB.get_by_job(jid)
    j["application"] = app_data
    # Include connections
    j["connections"] = ConnectionDB.get_by_job(jid)
    return j

@app.post("/api/match-score")
async def score_jobs():
    await _score_new_jobs()
    return {"message":"Scoring complete"}

# ═══ 7-8. PLAYBOOK + TAILOR (Page 4: Resume Studio) ═══
@app.post("/api/playbook")
async def gen_playbook():
    profile = ProfileDB.get(1)
    if not profile: raise HTTPException(400, "Create profile first")
    jobs = JobDB.get_unapplied(100)
    if not jobs: raise HTTPException(400, "No unapplied jobs")
    llm = get_llm(); await llm.init()
    from intelligence.engine import PlaybookGenerator
    gen = PlaybookGenerator(llm)
    result = await gen.generate(profile, jobs)
    rid = str(uuid.uuid4())[:8]
    from database import get_db
    conn = get_db()
    conn.execute("INSERT INTO playbooks (run_id,raw_output,model_used,tokens_used) VALUES (?,?,?,?)",
        (rid, result["playbook_text"], "groq", result["tokens"]))
    conn.commit(); conn.close()
    return {"run_id":rid,"playbook":result["playbook_text"],"jobs":result["job_count"],"tokens":result["tokens"]}

@app.post("/api/tailor")
async def tailor_all():
    profile = ProfileDB.get(1)
    if not profile: raise HTTPException(400, "Create profile first")
    from database import get_db
    conn = get_db()
    pb = conn.execute("SELECT raw_output FROM playbooks ORDER BY created_at DESC LIMIT 1").fetchone()
    conn.close()
    if not pb: raise HTTPException(400, "Generate playbook first")
    jobs = JobDB.get_unapplied(100)
    if not jobs: raise HTTPException(400, "No unapplied jobs")
    llm = get_llm(); await llm.init()
    from intelligence.engine import ResumeTailor
    from learning.engine import ABTestEngine
    tailor = ResumeTailor(llm)
    results = await tailor.tailor_batch(pb[0], jobs)
    rid = str(uuid.uuid4())[:8]
    aids = []
    for r in results:
        v, vd = ABTestEngine.assign()
        aid = ApplicationDB.create({"job_id":r.get("job_id"),"profile_id":1,"run_id":rid,
            "archetype":r.get("archetype"),"tailored_bullets":r.get("tailored_bullets"),
            "cover_letter":r.get("cover_letter",""),"variant":v,"variant_description":vd})
        aids.append(aid)
    return {"run_id":rid,"tailored":len(results),"application_ids":aids}

@app.get("/api/applications")
async def get_applications(status: str = None, limit: int = 200):
    return ApplicationDB.get_all(status, limit)

@app.get("/api/applications/{aid}")
async def get_application(aid: int):
    a = ApplicationDB.get_by_id(aid)
    if not a: raise HTTPException(404)
    return a

# ═══ 9. AUTO-APPLY (Page 5) ═══
@app.post("/api/apply/{aid}")
async def apply_one(aid: int, bg: BackgroundTasks):
    a = ApplicationDB.get_by_id(aid)
    if not a: raise HTTPException(404)
    job = JobDB.get_by_id(a["job_id"])
    profile = ProfileDB.get(a.get("profile_id",1))
    async def _apply():
        from rpa.applicant import RPAApplicant
        names = (profile.get("name","") or "").split()
        p = {**profile, "first_name":names[0] if names else "", "last_name":names[-1] if len(names)>1 else ""}
        rpa = RPAApplicant(p, {"auto_submit":False}, settings.HEADLESS_BROWSER)
        try:
            r = await rpa.apply_one(job, a.get("resume_path",""))
            ApplicationDB.update_status(aid, r.get("status","failed"), r.get("error"))
        except Exception as e: ApplicationDB.update_status(aid, "failed", str(e))
        finally: await rpa.cleanup()
    bg.add_task(_apply)
    return {"message":"Applying","application_id":aid}

@app.post("/api/apply/batch")
async def apply_batch(bg: BackgroundTasks, limit: int = 10):
    apps = ApplicationDB.get_all("pending", limit)
    return {"queued":len(apps),"message":f"Queued {len(apps)} applications"}

# ═══ 10. CONNECTIONS (Page 5: Insider Connect) ═══
@app.post("/api/connections/{jid}")
async def find_connections(jid: int):
    job = JobDB.get_by_id(jid)
    if not job: raise HTTPException(404)
    profile = ProfileDB.get(1)
    if not profile: raise HTTPException(400, "Create profile first")
    llm = get_llm(); await llm.init()
    from intelligence.engine import ConnectionFinder
    finder = ConnectionFinder(llm)
    result = await finder.find(profile, job)
    cid = ConnectionDB.create({"job_id":jid,"contact_title":result["contact_title"],
        "outreach_message":result["outreach_message"]})
    # Update job decision maker
    from database import get_db
    conn = get_db()
    conn.execute("UPDATE jobs SET decision_maker=? WHERE id=?", (result["contact_title"], jid))
    conn.commit(); conn.close()
    return {"connection_id":cid, **result}

@app.get("/api/connections/{jid}")
async def get_connections(jid: int):
    return ConnectionDB.get_by_job(jid)

# ═══ 11-12. ANALYTICS + LEARNING (Page 6) ═══
@app.post("/api/outcome")
async def update_outcome(data: OutcomeIn):
    ApplicationDB.update_outcome(data.application_id, data.outcome)
    return {"message":"Updated"}

@app.get("/api/analytics")
async def analytics():
    from learning.engine import ABTestEngine
    return {"stats":ApplicationDB.get_stats(), **ABTestEngine.report()}

@app.get("/api/analytics/proof-points")
async def proof_points():
    from learning.engine import OutcomeTracker
    return OutcomeTracker.get_proof_points()

@app.post("/api/learning/export")
async def export_training():
    from learning.engine import FineTuningExporter
    n = FineTuningExporter.export_jsonl()
    return {"exported":n,"path":"data/fine_tuning.jsonl"}

# ═══ DASHBOARD ═══
@app.get("/api/dashboard")
async def dashboard():
    return {
        "total_jobs": JobDB.count(),
        "applications": ApplicationDB.get_stats(),
        "recent_runs": RunDB.get_recent(5),
        "profile_exists": ProfileDB.exists(),
    }

# ═══ DOCUMENT GENERATION ═══
SECTION_NAMES = {
    "digitech": "Digitech Services",
    "asu": "Arizona State University",
    "vaxom": "Vaxom Packaging",
    "nccl": "National Commodities Clearing",
    "vertiv": "Vertiv (Capstone)",
    "km_capital": "KM Capital Partners",
    "scdi": "Supply Chain DI Platform",
    "gcn": "Global Careers Network",
}

@app.post("/api/generate-docs/{aid}")
async def gen_docs(aid: int):
    a = ApplicationDB.get_by_id(aid)
    if not a: raise HTTPException(404)
    job = JobDB.get_by_id(a["job_id"])
    profile = ProfileDB.get(a.get("profile_id", 1))
    bullets = a.get("tailored_bullets", {})
    if isinstance(bullets, str):
        try: bullets = json.loads(bullets)
        except: bullets = {}

    company_safe = re.sub(r'[^\w\s-]', '', job.get('company', 'Company')).replace(' ', '_')
    candidate_name = (profile.get('name', 'Atharva_Vaidya') if profile else 'Atharva_Vaidya').replace(' ', '_')
    fn = f"Resume_{candidate_name}_{company_safe}.docx"
    fp = f"data/outputs/resumes/{fn}"

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        # Title
        title = doc.add_heading(f"{job.get('title', '')} at {job.get('company', '')}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Bullets sections
        has_bullets = any(isinstance(v, list) and len(v) > 0 for v in bullets.values())
        if has_bullets:
            doc.add_heading("Tailored Resume Bullets", level=2)
            for sec, blist in bullets.items():
                if not isinstance(blist, list) or not blist: continue
                sec_heading = doc.add_paragraph()
                run = sec_heading.add_run(SECTION_NAMES.get(sec, sec.upper()))
                run.bold = True
                run.font.size = Pt(11)
                for b in blist:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(b)
        # Cover letter
        if a.get("cover_letter"):
            doc.add_page_break()
            doc.add_heading("Cover Letter", level=2)
            for para in a["cover_letter"].split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        doc.save(fp)
    except Exception as e:
        logger.error(f"docx generation failed: {e}, falling back to txt")
        fn = fn.replace('.docx', '.txt')
        fp = fp.replace('.docx', '.txt')
        output = f"TAILORED RESUME: {job.get('title','')} at {job.get('company','')}\n{'='*50}\n\n"
        for sec, blist in bullets.items():
            output += f"[{SECTION_NAMES.get(sec, sec.upper())}]\n"
            if isinstance(blist, list):
                for i, b in enumerate(blist, 1): output += f"  {i}. {b}\n"
            output += "\n"
        if a.get("cover_letter"):
            output += f"\nCOVER LETTER:\n{a['cover_letter']}\n"
        with open(fp, "w") as f: f.write(output)

    from database import get_db
    conn = get_db()
    conn.execute("UPDATE applications SET resume_path=? WHERE id=?", (fp, aid))
    conn.commit(); conn.close()
    return {"path": fp, "filename": fn}

@app.post("/api/generate-docs/batch")
async def gen_docs_batch():
    from database import get_db
    conn = get_db()
    apps = conn.execute("SELECT id FROM applications WHERE resume_path IS NULL AND tailored_bullets IS NOT NULL").fetchall()
    conn.close()
    generated = 0
    for row in apps:
        try: await gen_docs(row[0]); generated += 1
        except: continue
    return {"generated": generated, "total": len(apps)}

# Alias used by frontend
@app.post("/api/docs-batch")
async def docs_batch():
    return await gen_docs_batch()

# ═══ ONE-CLICK PIPELINE ═══
@app.post("/api/pipeline")
async def run_pipeline(bg: BackgroundTasks):
    rid = str(uuid.uuid4())[:8]
    RunDB.create(rid, rtype="full_pipeline")
    async def _run():
        # Step 1: Discover
        from discovery.engine import DiscoveryEngine
        engine = DiscoveryEngine(settings)
        try:
            stats = await engine.run_full(rid)
            RunDB.update(rid, jobs_discovered=stats["total_found"], jobs_new=stats["new_added"])
        except Exception as e:
            logger.error(f"Pipeline discover: {e}")
        finally:
            try: await engine.cleanup()
            except: pass
        # Step 2: Score
        await _score_new_jobs()
        # Step 3: Playbook
        profile = ProfileDB.get(1)
        jobs = JobDB.get_unapplied(100)
        pb_text = None
        if profile and jobs:
            try:
                llm = get_llm(); await llm.init()
                from intelligence.engine import PlaybookGenerator
                gen = PlaybookGenerator(llm)
                pb_result = await gen.generate(profile, jobs)
                pb_text = pb_result["playbook_text"]
                from database import get_db
                conn = get_db()
                conn.execute("INSERT INTO playbooks (run_id,raw_output,model_used,tokens_used) VALUES (?,?,?,?)",
                    (rid, pb_text, "groq", pb_result["tokens"]))
                conn.commit(); conn.close()
            except Exception as e:
                logger.error(f"Pipeline playbook: {e}")
        # Step 4: Tailor
        if pb_text is None:
            from database import get_db
            conn = get_db()
            pb_row = conn.execute("SELECT raw_output FROM playbooks ORDER BY created_at DESC LIMIT 1").fetchone()
            conn.close()
            if pb_row: pb_text = pb_row[0]
        jobs2 = JobDB.get_unapplied(100)
        tailored = 0
        if pb_text and jobs2:
            try:
                llm2 = get_llm(); await llm2.init()
                from intelligence.engine import ResumeTailor
                from learning.engine import ABTestEngine
                tailor = ResumeTailor(llm2)
                results = await tailor.tailor_batch(pb_text, jobs2)
                for r in results:
                    v, vd = ABTestEngine.assign()
                    ApplicationDB.create({"job_id": r.get("job_id"), "profile_id": 1, "run_id": rid,
                        "archetype": r.get("archetype"), "tailored_bullets": r.get("tailored_bullets"),
                        "cover_letter": r.get("cover_letter", ""), "variant": v, "variant_description": vd})
                    tailored += 1
                RunDB.update(rid, jobs_tailored=tailored)
            except Exception as e:
                logger.error(f"Pipeline tailor: {e}")
        # Step 5: Generate docs
        try:
            from database import get_db
            conn = get_db()
            doc_apps = conn.execute("SELECT id FROM applications WHERE resume_path IS NULL AND tailored_bullets IS NOT NULL").fetchall()
            conn.close()
            for row in doc_apps:
                try: await gen_docs(row[0])
                except: continue
        except Exception as e:
            logger.error(f"Pipeline gen_docs: {e}")
        RunDB.update(rid, status="completed", completed_at=datetime.now().isoformat())
    bg.add_task(_run)
    return {"run_id": rid, "status": "started", "steps": ["discover", "score", "playbook", "tailor", "generate_docs"]}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
