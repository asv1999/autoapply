"""Generate tailored resume and cover letter DOCX files in Atharva's template style."""
import json
import logging
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor, Twips

logger = logging.getLogger(__name__)

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
OUTPUT_DIR = os.path.join(BASE_DIR, "data", "outputs", "resumes")

BLACK = RGBColor(0x00, 0x00, 0x00)
NAVY = RGBColor(0x1B, 0x3A, 0x6B)
GRAY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

SZ_NAME = Pt(14)
SZ_CONTACT = Pt(8.5)
SZ_SECTION = Pt(9)
SZ_ROLE = Pt(9)
SZ_BODY = Pt(8.5)

WORK_SECTIONS = {
    "digitech": {
        "role": "Business Insights & Process Analytics Analyst",
        "org": "Digitech Services Limited",
        "location": "Phoenix, AZ",
        "dates": "Aug 2025 – Dec 2025",
    },
    "asu": {
        "role": "Analytics & Reporting Intern, Strategic Projects",
        "org": "Arizona State University",
        "location": "Tempe, AZ",
        "dates": "Aug 2024 – May 2025",
    },
    "vaxom": {
        "role": "Data Insights & Commercial Analytics Analyst",
        "org": "Vaxom Packaging",
        "location": "Mumbai, India",
        "dates": "Aug 2022 – Aug 2023",
    },
    "nccl": {
        "role": "Data Analytics & Reporting Analyst",
        "org": "National Commodities Clearing Limited",
        "location": "Mumbai, India",
        "dates": "Feb 2021 – Aug 2022",
    },
}

CONSULTING_SECTIONS = {
    "vertiv": {
        "role": "Performance Analytics & Benchmarking",
        "org": "Vertiv, Global Data Center Infrastructure",
        "location": "Thunderbird Capstone",
        "dates": "Sep 2024 – Dec 2024",
    },
    "km_capital": {
        "role": "Market Sizing & Data Analysis",
        "org": "KM Capital Partners, Healthcare Strategy",
        "location": "Thunderbird Corporate Partners",
        "dates": "Jan 2025 – May 2025",
    },
    "scdi": {
        "role": "Founder",
        "org": "Supply Chain Decision Intelligence Platform",
        "location": "Independent Project",
        "dates": "Jan 2026 – Present",
    },
}


def _slug(value, fallback):
    clean = re.sub(r"[^a-zA-Z0-9]+", "_", str(value or "")).strip("_")
    return clean or fallback


def _profile_name(profile):
    return profile.get("name") or "Atharva Vaidya"


def _name_slug(profile):
    return _slug(_profile_name(profile), "Candidate")


def _company_slug(job):
    return _slug(job.get("company"), "Unknown")


def _parse_jsonish(value, default):
    if value is None:
        return default
    if isinstance(value, str):
        try:
            return json.loads(value)
        except Exception:
            return default
    return value


def _normalize_bullets(application):
    bullets = _parse_jsonish(application.get("tailored_bullets"), {})
    return bullets if isinstance(bullets, dict) else {}


def _normalize_profile_bullets(profile):
    bullets = _parse_jsonish(profile.get("resume_bullets"), {})
    return bullets if isinstance(bullets, dict) else {}


def _has_resume_content(bullets):
    return any(
        isinstance(values, list) and any(len(str(item).strip()) > 20 for item in values)
        for values in bullets.values()
    )


def _set_run_style(run, *, size, color, bold=False, italic=False):
    run.font.name = "Calibri"
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _configure_resume_page(doc):
    for section in doc.sections:
        section.page_width = Twips(12240)
        section.page_height = Twips(15840)
        section.top_margin = Twips(720)
        section.bottom_margin = Twips(720)
        section.left_margin = Twips(900)
        section.right_margin = Twips(900)


def _configure_cover_letter_page(doc):
    for section in doc.sections:
        section.page_width = Twips(12240)
        section.page_height = Twips(15840)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _add_name(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    run = p.add_run((text or "").upper())
    _set_run_style(run, size=SZ_NAME, color=BLACK, bold=True)


def _add_contact(doc, values):
    parts = [str(value).strip() for value in values if str(value or "").strip()]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("  |  ".join(parts))
    _set_run_style(run, size=SZ_CONTACT, color=GRAY)


def _add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p_pr = p._element.get_or_add_pPr()
    p_bdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="1B3A6B"/>'
        "</w:pBdr>"
    )
    p_pr.append(p_bdr)


def _add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    _set_run_style(run, size=SZ_SECTION, color=NAVY, bold=True)


def _add_role_row(doc, role_text, org_text, location, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    p_pr = p._element.get_or_add_pPr()
    tabs = parse_xml(
        f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:pos="9300"/></w:tabs>'
    )
    p_pr.append(tabs)

    run = p.add_run(role_text)
    _set_run_style(run, size=SZ_ROLE, color=DARK, bold=True)

    location_part = f"  ·  {location}" if location else ""
    run = p.add_run(f"  ·  {org_text}{location_part}")
    _set_run_style(run, size=SZ_ROLE, color=DARK)

    run = p.add_run(f"\t{dates}")
    _set_run_style(run, size=SZ_CONTACT, color=GRAY, italic=True)


def _add_indented_paragraph(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1.5)
    p.paragraph_format.space_after = Pt(1.5)
    p_pr = p._element.get_or_add_pPr()
    p_pr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="300" w:hanging="180"/>'))
    return p


def _add_bullet(doc, text):
    p = _add_indented_paragraph(doc)
    run = p.add_run(f"•  {text}")
    _set_run_style(run, size=SZ_BODY, color=BLACK)


def _add_bold_bullet(doc, bold_part, rest):
    p = _add_indented_paragraph(doc)
    run = p.add_run("•  ")
    _set_run_style(run, size=SZ_BODY, color=BLACK)
    run = p.add_run(bold_part)
    _set_run_style(run, size=SZ_BODY, color=BLACK, bold=True)
    run = p.add_run(rest)
    _set_run_style(run, size=SZ_BODY, color=BLACK)


def _add_skill_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1.5)
    p.paragraph_format.space_after = Pt(1.5)
    run = p.add_run(f"{label}: ")
    _set_run_style(run, size=SZ_BODY, color=DARK, bold=True)
    run = p.add_run(value)
    _set_run_style(run, size=SZ_BODY, color=BLACK)


def _first_nonempty(*values):
    for value in values:
        if str(value or "").strip():
            return str(value).strip()
    return ""


def _linkedin(profile):
    raw = _first_nonempty(profile.get("linkedin_url"), profile.get("linkedin"))
    if not raw:
        return ""
    return raw.replace("https://", "").replace("http://", "").rstrip("/")


def _summary_bullets(profile, job):
    target_title = job.get("title") or "strategy and operations role"
    target_company = job.get("company") or "the team"
    sectors = "consulting, manufacturing, and technology"
    bullet_1 = (
        "4+ years",
        f" driving strategy, analytics, and process improvement across {sectors}"
        f" — aligning data, stakeholders, and execution for {target_title.lower()} work.",
    )
    bullet_2 = (
        "Track record:",
        " 115% revenue turnaround ($4M to $8.6M), 51% operational efficiency gains "
        "through AI transformation, and $100K+ in annual savings from workflow redesign.",
    )
    bullet_3 = (
        "Multi-market lens across 5 continents",
        f" — bringing structured problem-solving and cross-cultural execution to {target_company}.",
    )
    bullet_4 = (
        "Tools:",
        " SQL, Python, Excel, Tableau, Alteryx, Power BI, Salesforce, SAP EWM, Oracle ERP.",
    )
    return [bullet_1, bullet_2, bullet_3, bullet_4]


def _categorize_skills(skills):
    skills = _parse_jsonish(skills, skills)
    if isinstance(skills, str):
        skills = [part.strip() for part in skills.split(",") if part.strip()]
    if not isinstance(skills, list):
        skills = []

    analytics = []
    strategy = []
    tech = []

    analytics_kw = [
        "sql", "python", "tableau", "excel", "power bi", "looker", "alteryx",
        "data", "analytics", "bigquery", "google analytics",
    ]
    strategy_kw = [
        "strategy", "six sigma", "process", "transformation", "change management",
        "stakeholder", "business case", "financial", "market analysis", "gtm",
        "cross-functional", "leadership", "consulting", "client",
    ]
    tech_kw = [
        "salesforce", "sap", "crm", "automation", "ai", "digital twin",
        "scenario", "vba", "generative", "workflow", "oracle",
    ]

    for skill in skills:
        text = str(skill).strip()
        lower = text.lower()
        if any(token in lower for token in analytics_kw):
            analytics.append(text)
        elif any(token in lower for token in strategy_kw):
            strategy.append(text)
        elif any(token in lower for token in tech_kw):
            tech.append(text)
        elif text:
            analytics.append(text)

    return {
        "Analytics & Reporting": ", ".join(analytics) if analytics else "SQL, Python, Tableau, Excel, Power BI",
        "Strategy & Operations": ", ".join(strategy) if strategy else "Six Sigma DMAIC, Strategic Planning, Process Improvement",
        "Technology & Platforms": ", ".join(tech) if tech else "Salesforce CRM, SAP EWM, Oracle ERP",
    }


def _iter_section_bullets(section_name, tailored, base):
    for source in (tailored, base):
        values = source.get(section_name, [])
        if isinstance(values, list):
            cleaned = [str(item).strip() for item in values if len(str(item).strip()) > 20]
            if cleaned:
                return cleaned
    return []


def generate_resume(profile, application, job, output_dir=None):
    output_dir = output_dir or OUTPUT_DIR
    os.makedirs(output_dir, exist_ok=True)

    tailored_bullets = _normalize_bullets(application)
    base_bullets = _normalize_profile_bullets(profile)

    if not _has_resume_content(tailored_bullets) and not _has_resume_content(base_bullets):
        logger.warning("Resume generation skipped due to missing bullet content for application %s", application.get("id"))
        return None, None

    doc = Document()
    _configure_resume_page(doc)

    _add_name(doc, _profile_name(profile))
    _add_contact(
        doc,
        [
            _first_nonempty(profile.get("location"), "Phoenix, AZ"),
            profile.get("email"),
            profile.get("phone"),
            _linkedin(profile),
        ],
    )

    _add_divider(doc)
    _add_section_header(doc, "Summary")
    for bold_part, rest in _summary_bullets(profile, job):
        _add_bold_bullet(doc, bold_part, rest)

    _add_divider(doc)
    _add_section_header(doc, "Core Skills")
    for label, value in _categorize_skills(profile.get("skills", [])).items():
        _add_skill_line(doc, label, value)

    _add_divider(doc)
    _add_section_header(doc, "Professional Experience")
    for section_name in ("digitech", "asu", "vaxom", "nccl"):
        info = WORK_SECTIONS[section_name]
        _add_role_row(doc, info["role"], info["org"], info["location"], info["dates"])
        for bullet in _iter_section_bullets(section_name, tailored_bullets, base_bullets):
            _add_bullet(doc, bullet)

    _add_divider(doc)
    _add_section_header(doc, "Consulting Projects")
    for section_name in ("vertiv", "km_capital", "scdi"):
        info = CONSULTING_SECTIONS[section_name]
        _add_role_row(doc, info["role"], info["org"], info["location"], info["dates"])
        for bullet in _iter_section_bullets(section_name, tailored_bullets, base_bullets):
            _add_bullet(doc, bullet)

    _add_divider(doc)
    _add_section_header(doc, "Education")
    _add_role_row(
        doc,
        "Master of Global Management",
        "Thunderbird School of Global Management, ASU",
        "Phoenix, AZ",
        "May 2025",
    )
    _add_bullet(doc, "Thunderbird Alumni Scholarship (60% tuition)  |  GPA: 3.49/4.0")
    _add_role_row(
        doc,
        "B.Tech, Electronics & Telecommunications",
        "Vishwakarma Institute of Technology  ·  Research Scholar, KIST Seoul",
        "India",
        "Jun 2021",
    )

    _add_divider(doc)
    _add_section_header(doc, "Leadership & Interests")
    leadership_bullets = _iter_section_bullets("gcn", tailored_bullets, base_bullets)
    if leadership_bullets:
        for bullet in leadership_bullets:
            _add_bullet(doc, bullet)
    else:
        _add_bullet(
            doc,
            "President, Global Careers Network ASU (2024-25): Drove 12% YoY internship growth across "
            "14,000+ students by connecting them with Fortune 500 partners including JP Morgan Chase, "
            "Bain & Company, and McKinsey.",
        )
    _add_bullet(
        doc,
        'Cricket (AZ Premier League top run-scorer, 2 seasons)  |  Cooking (voted best vegetarian chef, '
        'Thunderbird 2025)  |  Newsletter: "Simplified by Atharva" — AI transformation in global operations',
    )

    filename = f"Resume_{_name_slug(profile)}_{_company_slug(job)}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    logger.info("Generated resume: %s", filepath)
    return filepath, filename


def generate_cover_letter(profile, application, job, output_dir=None):
    output_dir = output_dir or OUTPUT_DIR
    os.makedirs(output_dir, exist_ok=True)

    cover_letter = str(application.get("cover_letter") or "").strip()
    if len(cover_letter) < 50:
        return None, None

    doc = Document()
    _configure_cover_letter_page(doc)

    _add_name(doc, _profile_name(profile))
    _add_contact(
        doc,
        [
            profile.get("location"),
            profile.get("email"),
            profile.get("phone"),
            _linkedin(profile),
        ],
    )
    _add_divider(doc)
    doc.add_paragraph()

    for paragraph in re.split(r"\n\s*\n+", cover_letter):
        paragraph = paragraph.strip()
        if len(paragraph) < 20:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(paragraph)
        _set_run_style(run, size=Pt(10.5), color=BLACK)

    filename = f"CoverLetter_{_name_slug(profile)}_{_company_slug(job)}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    logger.info("Generated cover letter: %s", filepath)
    return filepath, filename
